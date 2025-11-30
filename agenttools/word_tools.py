"""Word document manipulation tools using python-docx."""

from typing import Annotated
from semantic_kernel.functions import kernel_function

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordTools:
    """Tools for creating and editing Word documents."""

    def __init__(self):
        """Initialize WordTools."""
        if not DOCX_AVAILABLE:
            print("Warning: python-docx not installed. Install with: pip install python-docx")

    @kernel_function(
        name="create_word_document",
        description="Create a new Word document at the specified path"
    )
    def create_word_document(
        self,
        file_path: Annotated[str, "Full path where the Word document will be created"]
    ) -> Annotated[str, "Success or error message"]:
        """Create a new blank Word document."""
        if not DOCX_AVAILABLE:
            return "Error: python-docx is not installed"
        
        try:
            doc = Document()
            doc.save(file_path)
            return f"Word document created successfully at: {file_path}"
        except Exception as err:
            return f"Error creating Word document: {str(err)}"

    @kernel_function(
        name="add_heading",
        description="Add a heading to a Word document"
    )
    def add_heading(
        self,
        file_path: Annotated[str, "Path to the Word document"],
        text: Annotated[str, "Heading text to add"],
        level: Annotated[int, "Heading level (0-9, where 0 is title)"] = 1
    ) -> Annotated[str, "Success or error message"]:
        """Add a heading to an existing Word document."""
        if not DOCX_AVAILABLE:
            return "Error: python-docx is not installed"
        
        try:
            doc = Document(file_path)
            doc.add_heading(text, level=level)
            doc.save(file_path)
            return f"Heading added successfully to: {file_path}"
        except Exception as err:
            return f"Error adding heading: {str(err)}"

    @kernel_function(
        name="add_paragraph",
        description="Add a paragraph to a Word document"
    )
    def add_paragraph(
        self,
        file_path: Annotated[str, "Path to the Word document"],
        text: Annotated[str, "Paragraph text to add"]
    ) -> Annotated[str, "Success or error message"]:
        """Add a paragraph to an existing Word document."""
        if not DOCX_AVAILABLE:
            return "Error: python-docx is not installed"
        
        try:
            doc = Document(file_path)
            doc.add_paragraph(text)
            doc.save(file_path)
            return f"Paragraph added successfully to: {file_path}"
        except Exception as err:
            return f"Error adding paragraph: {str(err)}"

    @kernel_function(
        name="add_table",
        description="Add a table to a Word document with specified rows and columns"
    )
    def add_table(
        self,
        file_path: Annotated[str, "Path to the Word document"],
        rows: Annotated[int, "Number of rows"],
        cols: Annotated[int, "Number of columns"]
    ) -> Annotated[str, "Success or error message"]:
        """Add a table to an existing Word document."""
        if not DOCX_AVAILABLE:
            return "Error: python-docx is not installed"
        
        try:
            doc = Document(file_path)
            doc.add_table(rows=rows, cols=cols)
            doc.save(file_path)
            return f"Table ({rows}x{cols}) added successfully to: {file_path}"
        except Exception as err:
            return f"Error adding table: {str(err)}"

    @kernel_function(
        name="set_table_cell",
        description="Set text in a specific table cell"
    )
    def set_table_cell(
        self,
        file_path: Annotated[str, "Path to the Word document"],
        table_index: Annotated[int, "Index of the table (0-based)"],
        row: Annotated[int, "Row index (0-based)"],
        col: Annotated[int, "Column index (0-based)"],
        text: Annotated[str, "Text to set in the cell"]
    ) -> Annotated[str, "Success or error message"]:
        """Set text in a specific table cell."""
        if not DOCX_AVAILABLE:
            return "Error: python-docx is not installed"
        
        try:
            doc = Document(file_path)
            if table_index >= len(doc.tables):
                return f"Error: Table index {table_index} out of range"
            
            table = doc.tables[table_index]
            if row >= len(table.rows) or col >= len(table.columns):
                return f"Error: Cell position ({row}, {col}) out of range"
            
            table.rows[row].cells[col].text = text
            doc.save(file_path)
            return f"Cell ({row}, {col}) updated successfully in table {table_index}"
        except Exception as err:
            return f"Error setting table cell: {str(err)}"

    @kernel_function(
        name="add_page_break",
        description="Add a page break to a Word document"
    )
    def add_page_break(
        self,
        file_path: Annotated[str, "Path to the Word document"]
    ) -> Annotated[str, "Success or error message"]:
        """Add a page break to an existing Word document."""
        if not DOCX_AVAILABLE:
            return "Error: python-docx is not installed"
        
        try:
            doc = Document(file_path)
            doc.add_page_break()
            doc.save(file_path)
            return f"Page break added successfully to: {file_path}"
        except Exception as err:
            return f"Error adding page break: {str(err)}"

    @kernel_function(
        name="read_word_document",
        description="Read all text content from a Word document"
    )
    def read_word_document(
        self,
        file_path: Annotated[str, "Path to the Word document"]
    ) -> Annotated[str, "Document content or error message"]:
        """Read all text from a Word document."""
        if not DOCX_AVAILABLE:
            return "Error: python-docx is not installed"
        
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as err:
            return f"Error reading Word document: {str(err)}"

    @kernel_function(
        name="add_bullet_list",
        description="Add a bulleted list to a Word document"
    )
    def add_bullet_list(
        self,
        file_path: Annotated[str, "Path to the Word document"],
        items: Annotated[str, "Comma-separated list items"]
    ) -> Annotated[str, "Success or error message"]:
        """Add a bulleted list to an existing Word document."""
        if not DOCX_AVAILABLE:
            return "Error: python-docx is not installed"
        
        try:
            doc = Document(file_path)
            item_list = [item.strip() for item in items.split(',')]
            for item in item_list:
                doc.add_paragraph(item, style='List Bullet')
            doc.save(file_path)
            return f"Bullet list with {len(item_list)} items added successfully"
        except Exception as err:
            return f"Error adding bullet list: {str(err)}"
